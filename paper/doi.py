#
# This script finds DOIs for a list of academic references using the CrossRef API
# and saves the results to a .docx file.
#
# --- PRE-REQUISITES ---
# You must install the required Python libraries before running this script.
# Open your terminal or command prompt and run the following commands:
#
# pip install requests
# pip install python-docx
#
# --- HOW TO USE ---
# 1. Make sure you have installed the libraries above.
# 2. Save this code as a Python file (e.g., find_dois.py).
# 3. Run the script from your terminal: python find_dois.py
# 4. A file named 'references_with_dois.docx' will be created in the same directory.
#

import requests
from docx import Document
import time

def get_doi_from_crossref(reference_text: str) -> str:
    """
    Queries the CrossRef API to find the DOI for a given reference string.

    Args:
        reference_text: The full bibliographic reference as a string.

    Returns:
        The found DOI as a string, or a "Not Found" message.
    """
    base_url = "https://api.crossref.org/works"
    
    # It's polite to identify your script with a User-Agent when using an API.
    # You can replace the email with your own if you plan to use this extensively.
    headers = {
        'User-Agent': 'DOI-Finder/1.0 (mailto:user@example.com)'
    }
    
    params = {
        'query.bibliographic': reference_text,
        'rows': 1  # We only need the top, most likely match.
    }
    
    try:
        response = requests.get(base_url, params=params, headers=headers, timeout=10)
        response.raise_for_status()  # Raises an HTTPError for bad responses (4xx or 5xx)
        
        data = response.json()
        
        # Check if the API returned any results
        if data.get('status') == 'ok' and data['message']['items']:
            # Get the DOI from the first item in the results
            doi = data['message']['items'][0].get('DOI')
            if doi:
                return doi
            else:
                return "DOI not found in API result."
        else:
            return "DOI not found."
            
    except requests.exceptions.RequestException as e:
        return f"API Request Error: {e}"
    except (KeyError, IndexError, TypeError):
        return "Could not parse API response."
    except Exception as e:
        return f"An unexpected error occurred: {e}"

def main():
    """
    Main function to process references and create the DOCX file.
    """
    # List of references to find DOIs for
    references = [
        "Ferreira, F. R. T., & do Couto, L. M. (2025). Using deep learning on microscopic images for white blood cell detection and segmentation to assist in leukemia diagnosis. The Journal of Supercomputing, 81(2), 410.",
        "Patel, H., Patel, G., & Patel, A. (2024, August). Blood Cancer Detection Using Machine Learning Techniques. In International Conference on Mobile Radio Communications & 5G Networks (pp. 473-482). Singapore: Springer Nature Singapore.",
        "Rastogi, P., Khanna, K., & Singh, V. (2022). LeuFeatx: Deep learning–based feature extractor for the diagnosis of acute leukemia from microscopic images of peripheral blood smear. Computers in Biology and Medicine, 142, 105236.",
        "Hassan, B. A. R., Mohammed, A. H., Hallit, S., Malaeb, D., & Hosseini, H. (2025). Exploring the role of artificial intelligence in chemotherapy development, cancer diagnosis, and treatment: present achievements and future outlook. Frontiers in Oncology, 15, 1475893.",
        "El Houby, E. M. (2025). Acute lymphoblastic leukemia diagnosis using machine learning techniques based on selected features. Scientific Reports, 15(1), 28056.",
        "Oybek Kizi, R. F., Theodore Armand, T. P., & Kim, H. C. (2025). A review of deep learning techniques for leukemia cancer classification based on blood smear images. Applied Biosciences, 4(1), 9.",
        "Yan, G., Mingyang, G., Wei, S., Hongping, L., Liyuan, Q., Ailan, L., ... & Yan, Q. (2025). Diagnosis and typing of leukemia using a single peripheral blood cell through deep learning. Cancer Science, 116(2), 533-543.",
        "Magotra, S., & Agrawal, K. (2024, July). Blood Cancer Detection Using Machine Learning: A Review. In 2024 1st International Conference on Sustainable Computing and Integrated Communication in Changing Landscape of AI (ICSCAI) (pp. 1-6). IEEE.",
        "Saranyan, N., Kanthimathi, N., Ramya, P., Kowsalya, N., & Mohanapriya, S. (2021, December). Blood cancer detection using machine learning. In 2021 5th international conference on electronics, communication and aerospace technology (ICECA) (pp. 1-11). IEEE.",
        "Karim, A., Azhari, A., Shahroz, M., Brahim Belhaouri, S., & Mustofa, K. (2022). LDSVM: Leukemia cancer classification using machine learning.",
        "Dharani, N. P., Sujatha, G., & Rani, R. (2023, August). Blood Cancer Detection Using Improved Machine Learning Algorithm. In 2023 International Conference on Circuit Power and Computing Technologies (ICCPCT) (pp. 1136-1141). IEEE.",
        "Ananth, C., Tamilselvi, P., Joshy, S. A., & Kumar, T. A. (2022). Blood Cancer Detection with Microscopic Images Using Machine Learning. In Machine Learning in Information and Communication Technology: Proceedings of ICICT 2021, SMIT (pp. 45-54). Singapore: Springer Nature Singapore.",
        "Malik, S., Iftikhar, A., Tauqeer, F. H., Adil, M., & Ahmed, S. (2022). A systematic literature review on leukemia prediction using machine learning. Journal of Computing & Biomedical Informatics, 3(02), 104-123.",
        "Bukhari, M., Yasmin, S., Sammad, S., & Abd El-Latif, A. A. (2022). A deep learning framework for leukemia cancer detection in microscopic blood samples using squeeze and excitation learning. Mathematical Problems in Engineering, 2022(1), 2801227.",
        "Awotunde, J. B., Imoize, A. L., Ayoade, O. B., Abiodun, M. K., Do, D. T., Silva, A., & Sur, S. N. (2022). An enhanced hyper-parameter optimization of a convolutional neural network model for leukemia cancer diagnosis in a smart healthcare system. Sensors, 22(24), 9689.",
        "Sallam, N. M., Saleh, A. I., Arafat Ali, H., & Abdelsalam, M. M. (2022). An efficient strategy for blood diseases detection based on grey wolf optimization as feature selection and machine learning techniques. Applied Sciences, 12(21), 10760.",
        "Tripti R, K., Bharathi, G., & Aditi, J. (2025). Comparison of Implementation in Blood Cancer Causes and Diseases. International Journal of Trend in Scientific Research and Development, 9(1), 503-512.",
        "Rupapara, V., Rustam, F., Aljedaani, W., Shahzad, H. F., Lee, E., & Ashraf, I. (2022). Blood cancer prediction using leukemia microarray gene data and hybrid logistic vector trees model. Scientific reports, 12(1), 1000.",
        "Iswarya, M. (2022, May). Detection of Leukemia using machine learning. In 2022 International Conference on Applied Artificial Intelligence and Computing (ICAAIC) (pp. 466-470). IEEE.",
        "Gupta, K., Jiwani, N., & Whig, P. (2022, September). Effectiveness of machine learning in detecting early-stage leukemia. In International Conference on Innovative Computing and Communications: Proceedings of ICICC 2022, Volume 2 (pp. 461-472). Singapore: Springer Nature Singapore.",
        "Švecová, M., Blahova, L., Kostolný, J., Birkova, A., Urdzik, P., Marekova, M., & Dubayova, K. (2025). Enhancing endometrial cancer detection: Blood serum intrinsic fluorescence data processing and machine learning application. Talanta, 283, 127083.",
        "Muthumanjula, M., & Bhoopalan, R. (2022). Detection of white blood cell cancer using deep learning using CMYK-moment localisation for information retrieval. Journal of IoT in Social, Mobile, Analytics, and Cloud, 4(1), 54-72.",
        "Das, P. K., Nayak, B., & Meher, S. (2022). A lightweight deep learning system for automatic detection of blood cancer. Measurement, 191, 110762.",
        "Ikechukwu, A. V., & Murali, S. (2022). i-Net: a deep CNN model for white blood cancer segmentation and classification. International Journal of Advanced Technology and Engineering Exploration, 9(95), 1448-1464.",
        "Amethiya, Y., Pipariya, P., Patel, S., & Shah, M. (2022). Comparative analysis of breast cancer detection using machine learning and biosensors. Intelligent Medicine, 2(2), 69-81.",
        "Chaturvedi, A., Pathak, N., Sharma, N., & Malik, P. (2025). Theoretical advancements in deep learning for hematological cancer diagnosis: proposing a new framework and mathematical model. Intelligent Decision Technologies, 19(3), 1893-1911.",
        "Alhumrani, S. Q., Ball, G. R., El-Sherif, A. A., Ahmed, S., Mousa, N. O., Alghorayed, S. A., ... & Gabre, R. M. (2025). Machine Learning for Multi-Omics Characterization of Blood Cancers: A Systematic Review. Cells, 14(17), 1385.",
        "Ananth, C., Mystica, K., Sridharan, M., James, S. A., & Kumar, T. A. (2022). An Advanced Low-cost Blood Cancer Detection System. International Journal of Early Childhood Special Education, 14(5).",
        "Baig, R., Rehman, A., Almuhaimeed, A., Alzahrani, A., & Rauf, H. T. (2022). Detecting malignant leukemia cells using microscopic blood smear images: a deep learning approach. Applied Sciences, 12(13), 6317.",
        "Shehta, A. I., Nasr, M., & El Ghazali, A. E. D. M. (2025). Blood cancer prediction model based on deep learning technique. Scientific Reports, 15(1), 1889."
    ]
    
    # Initialize a new Word document
    doc = Document()
    doc.add_heading('References with DOIs', level=1)
    
    total_refs = len(references)
    print(f"Starting to process {total_refs} references...")

    for i, ref_text in enumerate(references):
        print(f"[{i+1}/{total_refs}] Querying for: \"{ref_text[:70]}...\"")
        
        # Clean up reference string by removing extra whitespace
        cleaned_ref = ' '.join(ref_text.strip().split())
        
        # Get the DOI from the API
        doi_result = get_doi_from_crossref(cleaned_ref)
        
        # Add the numbered reference to the document
        p = doc.add_paragraph()
        p.add_run(f"{i+1}. ").bold = True
        p.add_run(f"{cleaned_ref} ")
        
        # Add the found DOI at the end, making it bold for emphasis
        p.add_run("DOI: ").bold = True
        p.add_run(doi_result).bold = True
        
        # Be polite to the API server by adding a small delay between requests
        time.sleep(0.5)

    # Save the final document
    output_filename = "references_with_dois.docx"
    try:
        doc.save(output_filename)
        print(f"\n✅ Success! All references processed.")
        print(f"File '{output_filename}' has been created in the current directory.")
    except Exception as e:
        print(f"\n❌ Error: Could not save the document. {e}")


if __name__ == "__main__":
    main()